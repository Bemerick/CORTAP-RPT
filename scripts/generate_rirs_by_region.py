#!/usr/bin/env python3
"""
Generate RIR Documents by Region from Excel Spreadsheet

This script generates RIR documents from the CORTAP Package 4 spreadsheet,
creating separate outputs for Region 1 (TRO-1) and Region 3 (TRO-3).

Usage:
    python scripts/generate_rirs_by_region.py

Input:
    docs/RIR 2026/RIR Cover Letter/CORTAP Package 4 - Reviews - SM Updated 123025.xlsx

Output:
    Generated documents saved to:
    - output/rir-documents-fy2026/{timestamp}/region1/
    - output/rir-documents-fy2026/{timestamp}/region3/
"""

import pandas as pd
import asyncio
from pathlib import Path
from datetime import datetime
import re
from html import escape
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Add parent directory to path for imports
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.services.document_generator import DocumentGenerator
from app.services.context_builder import RIRContextBuilder
from app.utils.logging import get_logger

logger = get_logger(__name__)


def sanitize_filename(name: str) -> str:
    """
    Sanitize a string for use in a filename.

    Args:
        name: String to sanitize

    Returns:
        Safe filename string
    """
    # Remove or replace invalid filename characters
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # Replace spaces with underscores
    name = name.replace(' ', '_')
    # Remove multiple underscores
    name = re.sub(r'_+', '_', name)
    # Limit length
    return name[:100]


def extract_region_number(region_code: str) -> int:
    """
    Extract region number from region code.

    Args:
        region_code: Region code (e.g., "TRO-1", "TRO-10")

    Returns:
        Region number as integer
    """
    # Extract number from region code (e.g., "TRO-1" -> 1)
    match = re.search(r'-?(\d+)$', str(region_code))
    if match:
        return int(match.group(1))
    # Default to 1 if can't extract
    return 1


def map_review_type(review_code: str) -> str:
    """
    Map review type code to full name.

    Args:
        review_code: Review code (e.g., "TR", "SMR", "Combined")

    Returns:
        Full review type name
    """
    mapping = {
        "TR": "Triennial Review",
        "TRIENNIAL": "Triennial Review",
        "TRIENNIAL REVIEW": "Triennial Review",
        "SMR": "State Management Review",
        "STATE MANAGEMENT REVIEW": "State Management Review",
        "COMBINED": "Combined Triennial and State Management Review",
        "COMBINED TRIENNIAL AND STATE MANAGEMENT REVIEW": "Combined Triennial and State Management Review"
    }
    return mapping.get(str(review_code).upper().strip(), "Triennial Review")


def format_phone(phone) -> str:
    """
    Format phone number or return TBD.

    Args:
        phone: Phone number value

    Returns:
        Formatted phone string
    """
    if pd.isna(phone):
        return "TBD"
    phone_str = str(phone).strip()
    if not phone_str or phone_str.upper() == "TBD":
        return "TBD"
    return phone_str


def format_date_for_display(date_value) -> str:
    """
    Format a date value for display in the RIR document.

    Args:
        date_value: Date value from Excel (datetime or string)

    Returns:
        Formatted date string (YYYY-MM-DD) or "TBD"
    """
    if pd.isna(date_value):
        return "TBD"

    if isinstance(date_value, datetime):
        return date_value.strftime("%Y-%m-%d")

    # Try to parse string dates
    try:
        dt = pd.to_datetime(date_value)
        return dt.strftime("%Y-%m-%d")
    except:
        return "TBD"


def add_hyperlink_to_text(paragraph, link_text: str, url: str):
    """
    Add a hyperlink to any text in a paragraph.

    This function finds and replaces plain text with a clickable hyperlink,
    preserving the original font while adding hyperlink styling.

    Args:
        paragraph: python-docx Paragraph object
        link_text: The text to find and make into a hyperlink
        url: The URL to link to
    """
    if not link_text or not url or link_text == "TBD" or link_text == "N/A":
        return

    from copy import deepcopy

    # Search for the link text in all runs
    for i, run in enumerate(paragraph.runs):
        if link_text in run.text:
            # Split the run text around the link text
            parts = run.text.split(link_text, 1)

            # Store reference to the original run element
            original_run_element = run._element

            # Update the current run to have text before link
            run.text = parts[0]

            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), '')

            # Set the hyperlink relationship
            r_id = paragraph.part.relate_to(
                url,
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink.set(qn('r:id'), r_id)

            # Create a new run for the hyperlink
            new_run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')

            # Explicitly set Times New Roman font for the hyperlink
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), 'Times New Roman')
            r_fonts.set(qn('w:hAnsi'), 'Times New Roman')
            r_fonts.set(qn('w:cs'), 'Times New Roman')
            r_pr.append(r_fonts)

            # Copy font size from original run if available
            if original_run_element.rPr is not None:
                r_sz = original_run_element.rPr.find(qn('w:sz'))
                if r_sz is not None:
                    r_pr.append(deepcopy(r_sz))

                r_sz_cs = original_run_element.rPr.find(qn('w:szCs'))
                if r_sz_cs is not None:
                    r_pr.append(deepcopy(r_sz_cs))

            # If no font size found, set default to 12pt (24 half-points)
            if original_run_element.rPr is None or original_run_element.rPr.find(qn('w:sz')) is None:
                r_sz = OxmlElement('w:sz')
                r_sz.set(qn('w:val'), '24')
                r_pr.append(r_sz)

                r_sz_cs = OxmlElement('w:szCs')
                r_sz_cs.set(qn('w:val'), '24')
                r_pr.append(r_sz_cs)

            # Add hyperlink color (blue)
            r_color = OxmlElement('w:color')
            r_color.set(qn('w:val'), '0563C1')  # Word hyperlink blue
            r_pr.append(r_color)

            # Add underline
            r_u = OxmlElement('w:u')
            r_u.set(qn('w:val'), 'single')
            r_pr.append(r_u)

            new_run.append(r_pr)

            # Add the link text with space preservation
            r_text = OxmlElement('w:t')
            r_text.set(qn('xml:space'), 'preserve')
            r_text.text = link_text
            new_run.append(r_text)

            hyperlink.append(new_run)

            # Insert the hyperlink right after the current run
            original_run_element.addnext(hyperlink)

            # Add remaining text if any - insert it right after the hyperlink
            if len(parts) > 1 and parts[1]:
                # Create a new run element for the remaining text
                remaining_run = OxmlElement('w:r')

                # Copy run properties from original using deepcopy
                if original_run_element.rPr is not None:
                    remaining_r_pr = deepcopy(original_run_element.rPr)
                    remaining_run.append(remaining_r_pr)

                # Add the remaining text with space preservation
                remaining_text = OxmlElement('w:t')
                remaining_text.set(qn('xml:space'), 'preserve')
                remaining_text.text = parts[1]
                remaining_run.append(remaining_text)

                # Insert right after the hyperlink
                hyperlink.addnext(remaining_run)

            return


def add_hyperlink_to_email(paragraph, email_text: str):
    """
    Add a mailto: hyperlink to an email address in a paragraph.

    This function finds and replaces plain text email with a clickable hyperlink,
    preserving the original font while adding hyperlink styling.

    Args:
        paragraph: python-docx Paragraph object
        email_text: The email address to find and hyperlink
    """
    if not email_text or email_text == "TBD":
        return

    from copy import deepcopy

    # Search for the email text in all runs
    for i, run in enumerate(paragraph.runs):
        if email_text in run.text:
            # Split the run text around the email
            parts = run.text.split(email_text, 1)

            # Store reference to the original run element
            original_run_element = run._element

            # Update the current run to have text before email
            run.text = parts[0]

            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), '')

            # Set the hyperlink relationship
            r_id = paragraph.part.relate_to(
                f'mailto:{email_text}',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink.set(qn('r:id'), r_id)

            # Create a new run for the hyperlink
            new_run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')

            # Explicitly set Times New Roman font for the hyperlink
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), 'Times New Roman')
            r_fonts.set(qn('w:hAnsi'), 'Times New Roman')
            r_fonts.set(qn('w:cs'), 'Times New Roman')
            r_pr.append(r_fonts)

            # Copy font size from original run if available
            if original_run_element.rPr is not None:
                r_sz = original_run_element.rPr.find(qn('w:sz'))
                if r_sz is not None:
                    r_pr.append(deepcopy(r_sz))

                r_sz_cs = original_run_element.rPr.find(qn('w:szCs'))
                if r_sz_cs is not None:
                    r_pr.append(deepcopy(r_sz_cs))

            # If no font size found, set default to 12pt (24 half-points)
            if original_run_element.rPr is None or original_run_element.rPr.find(qn('w:sz')) is None:
                r_sz = OxmlElement('w:sz')
                r_sz.set(qn('w:val'), '24')
                r_pr.append(r_sz)

                r_sz_cs = OxmlElement('w:szCs')
                r_sz_cs.set(qn('w:val'), '24')
                r_pr.append(r_sz_cs)

            # Add hyperlink color (blue)
            r_color = OxmlElement('w:color')
            r_color.set(qn('w:val'), '0563C1')  # Word hyperlink blue
            r_pr.append(r_color)

            # Add underline
            r_u = OxmlElement('w:u')
            r_u.set(qn('w:val'), 'single')
            r_pr.append(r_u)

            new_run.append(r_pr)

            # Add the email text with space preservation
            r_text = OxmlElement('w:t')
            r_text.set(qn('xml:space'), 'preserve')
            r_text.text = email_text
            new_run.append(r_text)

            hyperlink.append(new_run)

            # Insert the hyperlink right after the current run
            original_run_element.addnext(hyperlink)

            # Add remaining text if any - insert it right after the hyperlink
            if len(parts) > 1 and parts[1]:
                # Create a new run element for the remaining text
                remaining_run = OxmlElement('w:r')

                # Copy run properties from original using deepcopy
                if original_run_element.rPr is not None:
                    remaining_r_pr = deepcopy(original_run_element.rPr)
                    remaining_run.append(remaining_r_pr)

                # Add the remaining text with space preservation
                remaining_text = OxmlElement('w:t')
                remaining_text.set(qn('xml:space'), 'preserve')
                remaining_text.text = parts[1]
                remaining_run.append(remaining_text)

                # Insert right after the hyperlink
                hyperlink.addnext(remaining_run)

            return


def row_to_canonical_json(row: pd.Series) -> dict:
    """
    Convert a spreadsheet row to canonical JSON format.

    Args:
        row: Pandas Series representing one spreadsheet row

    Returns:
        Canonical JSON dict for RIR generation
    """
    # Extract recipient name and escape XML special characters
    recipient_name = escape(str(row.get('Recipient Name', '')).strip())

    # Format city, state
    city = str(row.get('City', '')).strip()
    state = str(row.get('State', '')).strip()
    recipient_city_state = f"{city}, {state}" if city and state else "N/A"

    # Extract region number
    region_number = extract_region_number(row.get('Region #', 'TRO-1'))

    # Map review type
    review_type = map_review_type(row.get('Type of Review', 'TR'))

    # Format site visit dates
    site_visit_start = format_date_for_display(row.get('FY26 Visit Start'))
    site_visit_end = format_date_for_display(row.get('FY26 Visit End (Complete by Sept 30, 2026)'))

    # Format due date
    due_date = format_date_for_display(row.get('FY26 RIR Due'))

    # Get FTA PM info
    fta_pm_name = str(row.get('FTA PM', '')).strip() or "TBD"
    fta_pm_title = str(row.get('FTA PM Title', '')).strip() or "TBD"
    fta_pm_phone = format_phone(row.get('FTA PM Phone'))
    fta_pm_email = str(row.get('FTA PM Email', '')).strip() or "TBD"

    # Get Lead Reviewer info
    lead_name = str(row.get('Lead', '')).strip() or "TBD"
    lead_email = str(row.get('Lead Email', '')).strip() or "TBD"
    lead_phone = format_phone(row.get('Lead Phone'))

    # Get website from spreadsheet
    website = str(row.get('Website', '')).strip() or "N/A"

    # Build canonical JSON structure
    canonical_json = {
        "project": {
            "region_number": region_number,
            "review_type": review_type,
            "recipient_name": recipient_name,
            "recipient_city_state": recipient_city_state,
            "recipient_id": str(row.get('Recipient ID', '')).strip(),
            "recipient_website": website,
            "site_visit_start_date": site_visit_start,
            "site_visit_end_date": site_visit_end,
            "due_date": due_date
        },
        "contractor": {
            "contractor_name": "Longevity Consulting",
            "lead_reviewer_name": lead_name,
            "lead_reviewer_phone": lead_phone,
            "lead_reviewer_email": lead_email
        },
        "fta_program_manager": {
            "fta_program_manager_name": fta_pm_name,
            "fta_program_manager_title": fta_pm_title,
            "fta_program_manager_phone": fta_pm_phone,
            "fta_program_manager_email": fta_pm_email
        }
    }

    return canonical_json


async def generate_rir_from_row(
    generator: DocumentGenerator,
    row: pd.Series,
    output_dir: Path,
    row_number: int
) -> tuple[bool, str]:
    """
    Generate a single RIR document from a spreadsheet row.

    Args:
        generator: DocumentGenerator instance
        row: Pandas Series with recipient data
        output_dir: Directory to save the document
        row_number: Row number for logging

    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        # Get recipient info for naming
        recipient_id = str(row.get('Recipient ID', '')).strip()
        recipient_name = str(row.get('Recipient Name', '')).strip()

        # Generate safe filename
        safe_name = sanitize_filename(recipient_name)
        filename = f"RIR_{recipient_id}_{safe_name}_FY2026.docx"
        output_path = output_dir / filename

        # Convert to canonical JSON
        json_data = row_to_canonical_json(row)

        # Build RIR context
        context = RIRContextBuilder.build_context(
            json_data,
            correlation_id=f"excel-{recipient_id}"
        )

        # Generate document
        document = await generator.generate(
            template_id="rir-package",
            context=context,
            correlation_id=f"excel-{recipient_id}"
        )

        # Save document
        with open(output_path, 'wb') as f:
            document.seek(0)
            f.write(document.read())

        # Post-process: Add hyperlinks to email addresses and website URLs
        # This approach is more Word-compatible than adding during rendering
        lead_email = json_data['contractor']['lead_reviewer_email']
        fta_pm_email = json_data['fta_program_manager']['fta_program_manager_email']
        recipient_website = json_data['project']['recipient_website']

        if (lead_email and lead_email != "TBD") or (fta_pm_email and fta_pm_email != "TBD") or (recipient_website and recipient_website not in ["N/A", "TBD"]):
            # Reopen document with python-docx
            docx_doc = Document(output_path)

            # Search all paragraphs for email addresses and website URLs, add hyperlinks
            for paragraph in docx_doc.paragraphs:
                if lead_email and lead_email != "TBD" and lead_email in paragraph.text:
                    add_hyperlink_to_email(paragraph, lead_email)
                if fta_pm_email and fta_pm_email != "TBD" and fta_pm_email in paragraph.text:
                    add_hyperlink_to_email(paragraph, fta_pm_email)
                if recipient_website and recipient_website not in ["N/A", "TBD"] and recipient_website in paragraph.text:
                    add_hyperlink_to_text(paragraph, recipient_website, recipient_website)

            # Also check tables (RIR documents may have tables)
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if lead_email and lead_email != "TBD" and lead_email in paragraph.text:
                                add_hyperlink_to_email(paragraph, lead_email)
                            if fta_pm_email and fta_pm_email != "TBD" and fta_pm_email in paragraph.text:
                                add_hyperlink_to_email(paragraph, fta_pm_email)
                            if recipient_website and recipient_website not in ["N/A", "TBD"] and recipient_website in paragraph.text:
                                add_hyperlink_to_text(paragraph, recipient_website, recipient_website)

            # Save the modified document
            docx_doc.save(output_path)

        file_size = output_path.stat().st_size
        message = f"✓ [{row_number:2d}] {recipient_name[:40]:40s} -> {filename} ({file_size:,} bytes)"
        return True, message

    except Exception as e:
        recipient_name = str(row.get('Recipient Name', 'Unknown'))[:40]
        message = f"✗ [{row_number:2d}] {recipient_name:40s} -> Error: {str(e)}"
        logger.error(f"Failed to generate RIR for row {row_number}", exc_info=True)
        return False, message


async def main():
    """Main script execution."""
    print("=" * 80)
    print("RIR Document Generator - By Region")
    print("=" * 80)

    # Setup paths
    project_root = Path(__file__).parent.parent
    excel_path = project_root / "docs" / "RIR 2026" / "RIR Cover Letter" / "CORTAP Package 4 - Reviews - SM Updated 123025.xlsx"
    excel_sheet = "Package 4 - Longevity"
    template_dir = project_root / "app" / "templates"

    # Create timestamped output directory
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_base = project_root / "output" / "rir-documents-fy2026" / timestamp
    output_region1 = output_base / "region1"
    output_region3 = output_base / "region3"

    # Verify Excel file exists
    if not excel_path.exists():
        print(f"\n✗ Error: Excel file not found: {excel_path}")
        sys.exit(1)

    # Create output directories
    output_region1.mkdir(parents=True, exist_ok=True)
    output_region3.mkdir(parents=True, exist_ok=True)

    print(f"\nInput file: {excel_path.name}")
    print(f"Sheet: {excel_sheet}")
    print(f"Output directories:")
    print(f"  Region 1: {output_region1}")
    print(f"  Region 3: {output_region3}")

    # Read Excel file
    print(f"\nReading Excel file...")
    try:
        # Read from Package 4 - Longevity sheet, skip first 2 rows (header row is row 3)
        df = pd.read_excel(excel_path, sheet_name=excel_sheet, skiprows=2)

        print(f"  Found {len(df)} recipients")

    except Exception as e:
        print(f"\n✗ Error reading Excel file: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Split by region
    region1_df = df[df['Region #'] == 'TRO-1'].reset_index(drop=True)
    region3_df = df[df['Region #'] == 'TRO-3'].reset_index(drop=True)

    print(f"\n  Region 1 (TRO-1): {len(region1_df)} recipients")
    print(f"  Region 3 (TRO-3): {len(region3_df)} recipients")

    # Initialize DocumentGenerator
    print(f"\nInitializing DocumentGenerator...")
    print(f"  Template directory: {template_dir}")
    generator = DocumentGenerator(template_dir=str(template_dir))

    # Generate Region 1 documents
    print(f"\n{'=' * 80}")
    print(f"Generating {len(region1_df)} Region 1 RIR documents...")
    print(f"{'=' * 80}\n")

    region1_results = []
    region1_messages = []

    for idx, row in region1_df.iterrows():
        success, message = await generate_rir_from_row(
            generator,
            row,
            output_region1,
            idx + 1
        )
        region1_results.append(success)
        region1_messages.append(message)
        print(message)

    # Generate Region 3 documents
    print(f"\n{'=' * 80}")
    print(f"Generating {len(region3_df)} Region 3 RIR documents...")
    print(f"{'=' * 80}\n")

    region3_results = []
    region3_messages = []

    for idx, row in region3_df.iterrows():
        success, message = await generate_rir_from_row(
            generator,
            row,
            output_region3,
            idx + 1
        )
        region3_results.append(success)
        region3_messages.append(message)
        print(message)

    # Summary
    print(f"\n{'=' * 80}")
    print(f"Generation Summary")
    print(f"{'=' * 80}")
    print(f"Region 1 (TRO-1):")
    print(f"  Total: {len(region1_df)}")
    print(f"  Successful: {sum(region1_results)}")
    print(f"  Failed: {len(region1_results) - sum(region1_results)}")

    print(f"\nRegion 3 (TRO-3):")
    print(f"  Total: {len(region3_df)}")
    print(f"  Successful: {sum(region3_results)}")
    print(f"  Failed: {len(region3_results) - sum(region3_results)}")

    print(f"\nOverall:")
    print(f"  Total: {len(df)}")
    print(f"  Successful: {sum(region1_results) + sum(region3_results)}")
    print(f"  Failed: {(len(region1_results) - sum(region1_results)) + (len(region3_results) - sum(region3_results))}")

    if all(region1_results) and all(region3_results):
        print(f"\n✓ All {len(df)} RIR documents generated successfully!")
        print(f"\nOutput directories:")
        print(f"  Region 1: {output_region1}")
        print(f"  Region 3: {output_region3}")
    else:
        print(f"\n⚠ Some documents failed to generate:")
        for success, message in zip(region1_results + region3_results, region1_messages + region3_messages):
            if not success:
                print(f"  {message}")


if __name__ == "__main__":
    # Run async main
    asyncio.run(main())
