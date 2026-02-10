#!/usr/bin/env python3
"""
Extract Merge Fields from Word Template

Purpose: Scan the draft report template and extract all text in square brackets [like this]
         to identify actual merge fields vs. instructional text.

Usage:
    python scripts/extract_template_fields.py

This helps identify what fields actually exist in the template for conversion.
"""

import re
from pathlib import Path
from docx import Document

def extract_fields_from_template(template_path: Path):
    """Extract all text in square brackets from Word template."""

    print(f"Analyzing template: {template_path.name}")
    print("=" * 80)

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template: {e}")
        return

    # Pattern to find text in square brackets
    # This will find [field], [Field Name], [FIELD], etc.
    bracket_pattern = re.compile(r'\[([^\]]+)\]')

    fields_found = set()
    instructional_text = set()

    # Search paragraphs
    print("\n📄 Scanning document paragraphs...")
    for para in doc.paragraphs:
        matches = bracket_pattern.findall(para.text)
        for match in matches:
            # Classify as field vs instruction
            if any(keyword in match.lower() for keyword in [
                'for triennial', 'for state management', 'for combined',
                'if the', 'delete', 'add as applicable', 'or', 'insert',
                'include the below', 'choose one'
            ]):
                instructional_text.add(match)
            else:
                fields_found.add(match)

    # Search tables
    print("📊 Scanning tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = bracket_pattern.findall(cell.text)
                for match in matches:
                    if any(keyword in match.lower() for keyword in [
                        'for triennial', 'for state management', 'delete',
                        'if the', 'add as applicable', 'or', 'insert'
                    ]):
                        instructional_text.add(match)
                    else:
                        fields_found.add(match)

    # Search headers/footers
    print("📋 Scanning headers and footers...")
    for section in doc.sections:
        # Header
        header = section.header
        matches = bracket_pattern.findall(header.paragraphs[0].text if header.paragraphs else "")
        for match in matches:
            fields_found.add(match)

        # Footer
        footer = section.footer
        matches = bracket_pattern.findall(footer.paragraphs[0].text if footer.paragraphs else "")
        for match in matches:
            fields_found.add(match)

    print("\n" + "=" * 80)
    print("RESULTS")
    print("=" * 80)

    # Display merge fields (likely actual data fields)
    print(f"\n✅ MERGE FIELDS FOUND ({len(fields_found)}):")
    print("-" * 80)
    if fields_found:
        for field in sorted(fields_found, key=str.lower):
            print(f"   [{field}]")
    else:
        print("   (None found)")

    # Display instructional text (needs to become conditional logic)
    print(f"\n📝 INSTRUCTIONAL TEXT ({len(instructional_text)}):")
    print("-" * 80)
    if instructional_text:
        for instruction in sorted(instructional_text, key=str.lower):
            print(f"   [{instruction}]")
    else:
        print("   (None found)")

    # Provide mapping suggestions
    print("\n" + "=" * 80)
    print("MAPPING SUGGESTIONS")
    print("=" * 80)
    print("\nBased on JSON structure, here's how to map common fields:\n")

    mapping_suggestions = {
        "Recipient name": "{{ project.recipient_name }}",
        "Recipient Name": "{{ project.recipient_name }}",
        "RECIPIENT NAME": "{{ project.recipient_name | upper }}",
        "Recipient Acronym": "{{ project.recipient_acronym }}",
        "RECIPIENT ACRONYM": "{{ project.recipient_acronym | upper }}",
        "Recipient ID": "{{ project.recipient_id }}",
        "City, State": "{{ project.recipient_city_state }}",
        "REGION #": "{{ project.region_number }}",
        "Region": "{{ project.region_number }}",
        "region": "{{ project.region_number }}",
        "fiscal year": "{{ project.fiscal_year }}",
        "FY": "{{ project.fiscal_year }}",
        "Month": "{{ project.report_date | date_format }}",
        "Day": "{{ project.report_date | date_format }}",
        "Year": "{{ project.fiscal_year }}",
        "FTA PM Name": "{{ fta_program_manager.name }}",
        "FTA Title": "{{ fta_program_manager.title }}",
        "phone number": "{{ fta_program_manager.phone }}",
        "email": "{{ fta_program_manager.email }}",
        "contractor name": "{{ contractor.lead_reviewer_name }}",
        "contractor firm": "{{ contractor.name }}",
        "reviewer name": "{{ contractor.lead_reviewer_name }}",
        "#": "{{ metadata.deficiency_count }}",
        "LIST": "{{ metadata.deficiency_areas }}",
    }

    for field in sorted(fields_found, key=str.lower):
        if field in mapping_suggestions:
            print(f"   [{field}]")
            print(f"      → {mapping_suggestions[field]}")
            print()

    # Warn about fields we DON'T see
    print("\n" + "=" * 80)
    print("⚠️  IMPORTANT NOTES")
    print("=" * 80)

    expected_dynamic_fields = ["review_type", "fiscal_year", "exit_conference_format"]

    print("\nSome JSON fields may NOT have direct template equivalents.")
    print("You may need to:")
    print("\n1. Add new dynamic fields where static text currently exists")
    print("   Example: If template says 'Triennial Review' in static text,")
    print("   replace with {{ project.review_type }}")

    print("\n2. Convert instructional text to conditional logic")
    print("   Example: '[For Triennial Reviews, delete...]'")
    print("   becomes: {% if project.review_type == 'Triennial Review' %}")

    print("\n3. Check for fields that might use different names")
    print("   Example: Template might use '[Date]' where JSON has 'report_date'")

    print("\n" + "=" * 80)
    print("\n💡 Next step: Open the template in Word and verify these fields")
    print("   Then use the conversion guide to map each to the JSON structure")
    print("\n")


def main():
    project_root = Path(__file__).parent.parent
    template_path = project_root / "app/templates/draft-audit-report-poc.docx"

    if not template_path.exists():
        print(f"❌ Template not found: {template_path}")
        print("\nExpected location: app/templates/draft-audit-report-poc.docx")
        return

    extract_fields_from_template(template_path)


if __name__ == '__main__':
    main()
