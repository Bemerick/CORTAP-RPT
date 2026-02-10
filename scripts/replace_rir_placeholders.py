#!/usr/bin/env python3
"""
Script to replace bracket placeholders in RIR template with Jinja2 syntax.

Replaces [Field Name] placeholders with {{ variable_name }} for python-docx template rendering.
"""

from docx import Document
import sys
import os


def replace_placeholders_in_paragraph(paragraph):
    """Replace bracket placeholders with Jinja2 syntax in a paragraph."""
    replacements = {
        '[#]': '{{ region_number }}',
        '[Triennial Review]': '{% if review_type == "Triennial Review" %}☒{% else %}☐{% endif %}',
        '[State Management Review]': '{% if review_type == "State Management Review" %}☒{% else %}☐{% endif %}',
        '[Combined Triennial and State Management Review]': '{% if review_type == "Combined Triennial and State Management Review" %}☒{% else %}☐{% endif %}',
        '[Recipient Name]': '{{ recipient_name }}',
        '[Recipient Location]': '{{ recipient_city_state }}',
        '[URL]': '{{ recipient_website }}',
        '[Contractor Name]': '{{ contractor_name }}',
        '[Lead Reviewer Name]': '{{ lead_reviewer_name }}',
        '[Lead Reviewer Phone #]': '{{ lead_reviewer_phone }}',
        '[Lead Reviewer Email Address]': '{{ lead_reviewer_email }}',
        '[FTA PM for Recipient]': '{{ fta_program_manager_name }}',
        '[FTA PM Title]': '{{ fta_program_manager_title }}',
        '[FTA PM Phone #]': '{{ fta_program_manager_phone }}',
        '[FTA PM Email Address]': '{{ fta_program_manager_email }}',
        # These are left as-is (Reserved for future use, Name/Service for recipient to fill)
        # '[Reserved]': '[Reserved]',
        # '[Name, Service Provided (Funding Source)]': '[Name, Service Provided (Funding Source)]',
    }

    text = paragraph.text
    modified = False

    for bracket, jinja in replacements.items():
        if bracket in text:
            text = text.replace(bracket, jinja)
            modified = True

    if modified:
        # Clear existing runs and add new text
        paragraph.clear()
        paragraph.add_run(text)
        return True

    return False


def replace_placeholders_in_table(table):
    """Replace bracket placeholders with Jinja2 syntax in a table."""
    replacements = {
        '[#]': '{{ region_number }}',
        '[Triennial Review]': '{% if review_type == "Triennial Review" %}☒{% else %}☐{% endif %}',
        '[State Management Review]': '{% if review_type == "State Management Review" %}☒{% else %}☐{% endif %}',
        '[Combined Triennial and State Management Review]': '{% if review_type == "Combined Triennial and State Management Review" %}☒{% else %}☐{% endif %}',
        '[Recipient Name]': '{{ recipient_name }}',
        '[Recipient Location]': '{{ recipient_city_state }}',
        '[URL]': '{{ recipient_website }}',
        '[Contractor Name]': '{{ contractor_name }}',
        '[Lead Reviewer Name]': '{{ lead_reviewer_name }}',
        '[Lead Reviewer Phone #]': '{{ lead_reviewer_phone }}',
        '[Lead Reviewer Email Address]': '{{ lead_reviewer_email }}',
        '[FTA PM for Recipient]': '{{ fta_program_manager_name }}',
        '[FTA PM Title]': '{{ fta_program_manager_title }}',
        '[FTA PM Phone #]': '{{ fta_program_manager_phone }}',
        '[FTA PM Email Address]': '{{ fta_program_manager_email }}',
    }

    modified_count = 0

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text
                modified = False

                for bracket, jinja in replacements.items():
                    if bracket in text:
                        text = text.replace(bracket, jinja)
                        modified = True

                if modified:
                    paragraph.clear()
                    paragraph.add_run(text)
                    modified_count += 1

    return modified_count


def main():
    input_file = 'app/templates/RO_State_Recipient_FY2026_RecipientInformationRequestPackageTemplate.docx'
    output_file = 'app/templates/RO_State_Recipient_FY2026_RecipientInformationRequestPackageTemplate_updated.docx'

    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    print(f"Reading template: {input_file}")
    doc = Document(input_file)

    # Replace in paragraphs
    para_count = 0
    for paragraph in doc.paragraphs:
        if replace_placeholders_in_paragraph(paragraph):
            para_count += 1

    # Replace in tables
    table_cell_count = 0
    for table in doc.tables:
        table_cell_count += replace_placeholders_in_table(table)

    # Save the modified document
    print(f"Saving updated template: {output_file}")
    doc.save(output_file)

    print(f"\n✓ Replacement complete!")
    print(f"  - Modified {para_count} paragraphs")
    print(f"  - Modified {table_cell_count} table cells")
    print(f"\nOutput saved to: {output_file}")
    print("\nNext steps:")
    print("  1. Review the updated template in Word")
    print("  2. If satisfied, rename it to replace the original template")
    print("  3. Test document generation with the new template")


if __name__ == '__main__':
    main()
