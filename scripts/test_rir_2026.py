#!/usr/bin/env python3
"""
Test script for 2026 RIR template generation.

Tests the new FY2026 RIR template with sample data.
"""

import asyncio
import sys
from pathlib import Path
from datetime import datetime

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from app.services.document_generator import DocumentGenerator
from app.services.context_builder import RIRContextBuilder
from app.models.template_data import RIRTemplateData


async def test_rir_2026_generation():
    """Test RIR generation with 2026 template."""

    print("=" * 80)
    print("Testing FY2026 RIR Template Generation")
    print("=" * 80)

    # Sample test data for Connecticut DOT
    test_data = RIRTemplateData(
        region_number=1,
        review_type="Triennial Review",
        recipient_name="Connecticut Department of Transportation",
        recipient_city_state="Newington, CT",
        recipient_id="1334",
        recipient_website="https://portal.ct.gov/DOT",
        site_visit_dates="March 10-14, 2026",
        contractor_name="Qi Tech, LLC",
        lead_reviewer_name="Sarah Johnson",
        lead_reviewer_phone="(617) 555-1234",
        lead_reviewer_email="sarah.johnson@qitech.com",
        fta_program_manager_name="Michael Chen",
        fta_program_manager_title="Transportation Program Manager",
        fta_program_manager_phone="(617) 494-3500",
        fta_program_manager_email="michael.chen@dot.gov",
        due_date="February 10, 2026"
    )

    print("\nTest Data:")
    print(f"  Region: {test_data.region_number}")
    print(f"  Review Type: {test_data.review_type}")
    print(f"  Recipient: {test_data.recipient_name}")
    print(f"  Location: {test_data.recipient_city_state}")
    print(f"  Recipient ID: {test_data.recipient_id}")
    print(f"  Contractor: {test_data.contractor_name}")
    print(f"  Lead Reviewer: {test_data.lead_reviewer_name}")

    # Initialize document generator
    template_dir = project_root / "app" / "templates"
    generator = DocumentGenerator(str(template_dir))

    print(f"\nTemplate Directory: {template_dir}")
    print(f"Using Template: rir-package.docx")

    # Convert to context dictionary
    context = test_data.to_template_context()

    print("\nGenerating document...")
    try:
        # Generate document
        document_bytes = await generator.generate(
            template_id="rir-package",
            context=context,
            correlation_id="test-2026-001"
        )

        # Save to output directory
        output_dir = project_root / "output" / "test-rir-2026"
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"RIR_Test_FY2026_{timestamp}.docx"

        with open(output_file, 'wb') as f:
            f.write(document_bytes.getvalue())

        print(f"\n✓ Document generated successfully!")
        print(f"  Output: {output_file}")
        print(f"  Size: {len(document_bytes.getvalue()):,} bytes")

        # Verify the document
        print("\nVerifying generated document...")
        from docx import Document
        doc = Document(output_file)

        # Check for Jinja2 patterns (should be replaced)
        text_content = "\n".join([para.text for para in doc.paragraphs])

        if "{{" in text_content or "{%" in text_content:
            print("  ⚠️  WARNING: Found unreplaced Jinja2 templates in output!")
            # Find and print them
            for para in doc.paragraphs:
                if "{{" in para.text or "{%" in para.text:
                    print(f"    - {para.text[:100]}")
        else:
            print("  ✓ No unreplaced Jinja2 templates found")

        # Check for populated fields
        print("\nField Population Check:")
        checks = [
            ("Region number", test_data.region_number, str(test_data.region_number)),
            ("Recipient name", test_data.recipient_name, test_data.recipient_name),
            ("Contractor", test_data.contractor_name, test_data.contractor_name),
            ("Lead Reviewer", test_data.lead_reviewer_name, test_data.lead_reviewer_name),
        ]

        for field_name, expected, search_text in checks:
            if search_text in text_content:
                print(f"  ✓ {field_name}: Found")
            else:
                print(f"  ✗ {field_name}: NOT FOUND (expected '{expected}')")

        print("\n" + "=" * 80)
        print("Test Complete!")
        print("=" * 80)

        return output_file

    except Exception as e:
        print(f"\n✗ Error generating document: {e}")
        import traceback
        traceback.print_exc()
        return None


if __name__ == "__main__":
    result = asyncio.run(test_rir_2026_generation())
    sys.exit(0 if result else 1)
