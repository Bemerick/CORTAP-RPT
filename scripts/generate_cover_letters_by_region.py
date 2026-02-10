#!/usr/bin/env python3
"""
Generate RIR Cover Letters by Region from Excel Spreadsheet

This script generates RIR cover letters from the CORTAP Package 4 spreadsheet,
using region-specific templates for Region 1 and Region 3. It handles 1-3
Recipient POCs in the cc: line, with each POC on a separate line.

Usage:
    python scripts/generate_cover_letters_by_region.py

Input:
    docs/RIR 2026/RIR Cover Letter/CORTAP Package 4 - Reviews - SM Updated 123025.xlsx

Output:
    Generated documents saved to:
    - output/cover-letters-fy2026/{timestamp}/region1/
    - output/cover-letters-fy2026/{timestamp}/region3/
"""

import pandas as pd
from pathlib import Path
from datetime import datetime
import re
from docxtpl import DocxTemplate
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Add parent directory to path for imports
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))


def sanitize_filename(name: str) -> str:
    """
    Sanitize a string for use in a filename.

    Args:
        name: String to sanitize

    Returns:
        Safe filename string
    """
    # Remove or replace invalid filename characters
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # Replace spaces with underscores
    name = name.replace(' ', '_')
    # Remove multiple underscores
    name = re.sub(r'_+', '_', name)
    # Limit length
    return name[:100]


def map_review_type(review_code: str) -> str:
    """
    Map review type code to full name.

    Args:
        review_code: Review code (e.g., "TR", "SMR", "Combined")

    Returns:
        Full review type name
    """
    mapping = {
        "TR": "Triennial Review",
        "TRIENNIAL": "Triennial Review",
        "TRIENNIAL REVIEW": "Triennial Review",
        "SMR": "State Management Review",
        "STATE MANAGEMENT REVIEW": "State Management Review",
        "COMBINED": "Combined Triennial and State Management Review",
        "COMBINED TRIENNIAL AND STATE MANAGEMENT REVIEW": "Combined Triennial and State Management Review"
    }
    return mapping.get(str(review_code).upper().strip(), "Triennial Review")


def add_article_to_review_type(review_type: str) -> str:
    """
    Add article (a/an) to review type.

    Args:
        review_type: Review type (e.g., "Triennial Review")

    Returns:
        Review type with article (e.g., "a Triennial Review")
    """
    review_type = str(review_type).strip()

    # Determine article based on first letter
    if review_type[0].lower() in ['a', 'e', 'i', 'o', 'u']:
        return f"an {review_type}"
    return f"a {review_type}"


def escape_xml_chars(text: str) -> str:
    """
    Escape XML special characters in text.

    Args:
        text: Text to escape

    Returns:
        XML-escaped text
    """
    if not text or text == "TBD":
        return text
    return xml_escape(str(text))


def format_date_for_letter(date_value) -> str:
    """
    Format a date value for display in the cover letter.

    Args:
        date_value: Date value from Excel (datetime or string)

    Returns:
        Formatted date string (e.g., "February 27, 2026") or "TBD"
    """
    if pd.isna(date_value):
        return "TBD"

    if isinstance(date_value, datetime):
        return date_value.strftime("%B %d, %Y")

    # Try to parse string dates
    try:
        dt = pd.to_datetime(date_value)
        return dt.strftime("%B %d, %Y")
    except:
        return "TBD"


def build_poc_cc_line(row: pd.Series) -> str:
    """
    Build the POC cc: line from 1-3 Recipient POCs.
    Each POC is on their own line with the recipient organization name.

    Args:
        row: Pandas Series with recipient data

    Returns:
        Formatted cc: line with each POC on separate line with organization
        Example: "Ms. Lisa Rivers, Connecticut Department of Transportation
    Mr. Rich Jankovich, Connecticut Department of Transportation"
    """
    poc_lines = []

    # Get recipient name (organization)
    recipient_name = escape_xml_chars(str(row.get('Recipient Name', '')).strip())

    # Check up to 3 POCs
    for poc_num in [1, 2, 3]:
        poc_salutation_raw = row.get(f'Recipient POC {poc_num} Name Salutation')
        poc_name_raw = row.get(f'Recipient POC {poc_num} Name')

        # Only include if name is present
        if pd.notna(poc_name_raw) and str(poc_name_raw).strip():
            poc_salutation = str(poc_salutation_raw).strip() if pd.notna(poc_salutation_raw) else ""
            poc_name = str(poc_name_raw).strip()

            # Format: "Salutation Name, Organization"
            if poc_salutation:
                poc_line = f"{poc_salutation} {poc_name}, {recipient_name}"
            else:
                poc_line = f"{poc_name}, {recipient_name}"

            poc_lines.append(poc_line)

    # Join with newline and tab (to align with template's indentation)
    return "\n\t".join(poc_lines)


def row_to_cover_letter_context(row: pd.Series) -> dict:
    """
    Convert a spreadsheet row to cover letter template context.

    Args:
        row: Pandas Series representing one spreadsheet row

    Returns:
        Context dict for cover letter template
    """
    # Get basic recipient info from Excel columns (escape XML chars)
    recipient_name = escape_xml_chars(str(row.get('Recipient Name', '')).strip())
    recipient_acronym = escape_xml_chars(str(row.get('Recipient Accronym', '')).strip())
    address = escape_xml_chars(str(row.get('Recipient Address', '')).strip())
    city = escape_xml_chars(str(row.get('Recipient City', '')).strip())
    state = str(row.get('Recipient State', '')).strip()

    # Handle zip code - ensure 5 digits with leading zeros
    zip_raw = row.get('Recipient Zip', '')
    if pd.notna(zip_raw):
        # Convert to string and remove any decimals (Excel may store as float)
        zip_str = str(int(float(zip_raw))) if zip_raw else ''
        # Pad with leading zeros to ensure 5 digits
        zip_code = zip_str.zfill(5) if zip_str else ''
    else:
        zip_code = ''

    # Get review type from Excel and map to full name
    review_type_code = str(row.get('Review Type', 'TR')).strip()
    review_type = map_review_type(review_type_code)
    review_type_article = add_article_to_review_type(review_type)

    # Get contractor name from Excel
    contractor_name = str(row.get('Contractor Company', 'Longevity')).strip()
    # If it's just "Longevity", expand to full name
    if contractor_name == "Longevity":
        contractor_name = "Longevity Consulting"

    # Get lead reviewer info
    lead_salutation = str(row.get('Lead Reviewer Salutation', 'Mr./Ms.')).strip()
    lead_name = str(row.get('Lead Reviewer Name', 'TBD')).strip()
    lead_title_name = f"{lead_salutation} {lead_name}" if lead_name != "TBD" else "TBD"

    # Get lead email and phone from Excel (use TBD if not available)
    lead_email_raw = row.get('Lead Reviewer Email')
    lead_email = str(lead_email_raw).strip() if pd.notna(lead_email_raw) and str(lead_email_raw).strip() else "TBD"

    lead_phone_raw = row.get('Lead Reviewer Phone')
    lead_phone = str(lead_phone_raw).strip() if pd.notna(lead_phone_raw) and str(lead_phone_raw).strip() else "TBD"

    # Get Regional Office Rep info
    ro_rep_salutation = str(row.get('Regional Office Representative Salutation', 'Mr./Ms.')).strip()
    ro_rep_name = str(row.get('Regional Office Representative Name', 'TBD')).strip()
    ro_rep_title_name = f"{ro_rep_salutation} {ro_rep_name}" if ro_rep_name != "TBD" else "TBD"
    ro_rep_phone = str(row.get('Regional Office Representative Phone', 'TBD')).strip()
    ro_rep_email = str(row.get('Regional Office Representative email', 'TBD')).strip()

    # Get Regional Administrator name
    regional_admin_name = str(row.get('Regional Office Administrator Name', 'Regional Administrator')).strip()

    # Get due date from Date column
    letter_date_value = row.get('Date')
    letter_date = format_date_for_letter(letter_date_value)

    # Due date is calculated as letter date + ~45 days (February 27, 2026)
    # For now, use fixed date from original template
    rir_due_date_formatted = "February 27, 2026"

    # Recipient contact info from Excel (escape XML chars)
    recipient_contact_name = escape_xml_chars(str(row.get('Recipient Accountable Executive Name', 'Transit Agency Contact')).strip())
    recipient_contact_title = escape_xml_chars(str(row.get('Recipient Accountable Executive Title', 'Executive Director')).strip())

    # Get salutation and last name for Region 3 greeting (columns L and N)
    recipient_salutation = escape_xml_chars(str(row.get('Recipient Accountable Executive Salutation', 'Mr./Ms.')).strip())
    recipient_last_name = escape_xml_chars(str(row.get('RAE Last Name', '')).strip())

    # Build POC cc: line
    recipient_poc_cc_line = build_poc_cc_line(row)

    # Build context
    context = {
        'letter_date': letter_date,
        'recipient_contact_name': recipient_contact_name,
        'recipient_contact_title': recipient_contact_title,
        'recipient_salutation': recipient_salutation,
        'recipient_last_name': recipient_last_name,
        'recipient_name': recipient_name,
        'street_address': address,
        'city': city,
        'state': state,
        'zip_code': zip_code,
        'review_type': review_type,
        'review_type_article': review_type_article,
        'recipient_acronym': recipient_acronym,
        'contractor_name': contractor_name,
        'lead_title_name': lead_title_name,
        'lead_email': lead_email,
        'lead_phone': lead_phone,
        'rir_due_date_formatted': rir_due_date_formatted,
        'ro_rep_title_name': ro_rep_title_name,
        'ro_rep_phone': ro_rep_phone,
        'ro_rep_email': ro_rep_email,
        'regional_admin_name': regional_admin_name,
        'recipient_poc_cc_line': recipient_poc_cc_line,
    }

    return context


def add_hyperlink_to_email(paragraph, email_text: str):
    """
    Add a mailto: hyperlink to an email address in a paragraph.

    This function finds and replaces plain text email with a clickable hyperlink,
    preserving the original font while adding hyperlink styling.

    Args:
        paragraph: python-docx Paragraph object
        email_text: The email address to find and hyperlink
    """
    if not email_text or email_text == "TBD":
        return

    from copy import deepcopy

    # Search for the email text in all runs
    for i, run in enumerate(paragraph.runs):
        if email_text in run.text:
            # Split the run text around the email
            parts = run.text.split(email_text, 1)

            # Store reference to the original run element
            original_run_element = run._element

            # Update the current run to have text before email
            run.text = parts[0]

            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), '')

            # Set the hyperlink relationship
            r_id = paragraph.part.relate_to(
                f'mailto:{email_text}',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                is_external=True
            )
            hyperlink.set(qn('r:id'), r_id)

            # Create a new run for the hyperlink
            new_run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')

            # Explicitly set Times New Roman font for the hyperlink
            r_fonts = OxmlElement('w:rFonts')
            r_fonts.set(qn('w:ascii'), 'Times New Roman')
            r_fonts.set(qn('w:hAnsi'), 'Times New Roman')
            r_fonts.set(qn('w:cs'), 'Times New Roman')
            r_pr.append(r_fonts)

            # Copy font size from original run if available
            if original_run_element.rPr is not None:
                r_sz = original_run_element.rPr.find(qn('w:sz'))
                if r_sz is not None:
                    r_pr.append(deepcopy(r_sz))

                r_sz_cs = original_run_element.rPr.find(qn('w:szCs'))
                if r_sz_cs is not None:
                    r_pr.append(deepcopy(r_sz_cs))

            # If no font size found, set default to 12pt (24 half-points)
            if original_run_element.rPr is None or original_run_element.rPr.find(qn('w:sz')) is None:
                r_sz = OxmlElement('w:sz')
                r_sz.set(qn('w:val'), '24')
                r_pr.append(r_sz)

                r_sz_cs = OxmlElement('w:szCs')
                r_sz_cs.set(qn('w:val'), '24')
                r_pr.append(r_sz_cs)

            # Add hyperlink color (blue)
            r_color = OxmlElement('w:color')
            r_color.set(qn('w:val'), '0563C1')  # Word hyperlink blue
            r_pr.append(r_color)

            # Add underline
            r_u = OxmlElement('w:u')
            r_u.set(qn('w:val'), 'single')
            r_pr.append(r_u)

            new_run.append(r_pr)

            # Add the email text with space preservation
            r_text = OxmlElement('w:t')
            r_text.set(qn('xml:space'), 'preserve')
            r_text.text = email_text
            new_run.append(r_text)

            hyperlink.append(new_run)

            # Insert the hyperlink right after the current run
            original_run_element.addnext(hyperlink)

            # Add remaining text if any - insert it right after the hyperlink
            if len(parts) > 1 and parts[1]:
                # Create a new run element for the remaining text
                remaining_run = OxmlElement('w:r')

                # Copy run properties from original using deepcopy
                if original_run_element.rPr is not None:
                    remaining_r_pr = deepcopy(original_run_element.rPr)
                    remaining_run.append(remaining_r_pr)

                # Add the remaining text with space preservation
                remaining_text = OxmlElement('w:t')
                remaining_text.set(qn('xml:space'), 'preserve')
                remaining_text.text = parts[1]
                remaining_run.append(remaining_text)

                # Insert right after the hyperlink
                hyperlink.addnext(remaining_run)

            return


def generate_cover_letter_from_row(
    template_path: Path,
    row: pd.Series,
    output_dir: Path,
    row_number: int
) -> tuple[bool, str]:
    """
    Generate a single cover letter document from a spreadsheet row.

    Args:
        template_path: Path to the Jinja2 template file
        row: Pandas Series with recipient data
        output_dir: Directory to save the document
        row_number: Row number for logging

    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        # Get recipient info for naming
        recipient_id = str(row.get('Recipient ID', '')).strip()
        recipient_name = str(row.get('Recipient Name', '')).strip()

        # Generate safe filename
        safe_name = sanitize_filename(recipient_name)
        filename = f"CoverLetter_{recipient_id}_{safe_name}_FY2026.docx"
        output_path = output_dir / filename

        # Load template first (needed for building hyperlinks)
        doc = DocxTemplate(template_path)

        # Build context with plain strings
        context = row_to_cover_letter_context(row)

        # Render template with context
        doc.render(context)

        # Save document
        doc.save(output_path)

        # Post-process: Add hyperlinks to email addresses
        # This approach is more Word-compatible than using RichText during rendering
        lead_email = context.get('lead_email')
        ro_rep_email = context.get('ro_rep_email')

        if (lead_email and lead_email != "TBD") or (ro_rep_email and ro_rep_email != "TBD"):
            # Reopen document with python-docx
            docx_doc = Document(output_path)

            # Search all paragraphs for email addresses and add hyperlinks
            for paragraph in docx_doc.paragraphs:
                if lead_email and lead_email != "TBD" and lead_email in paragraph.text:
                    add_hyperlink_to_email(paragraph, lead_email)
                if ro_rep_email and ro_rep_email != "TBD" and ro_rep_email in paragraph.text:
                    add_hyperlink_to_email(paragraph, ro_rep_email)

            # Also check tables (cover letters may have tables)
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if lead_email and lead_email != "TBD" and lead_email in paragraph.text:
                                add_hyperlink_to_email(paragraph, lead_email)
                            if ro_rep_email and ro_rep_email != "TBD" and ro_rep_email in paragraph.text:
                                add_hyperlink_to_email(paragraph, ro_rep_email)

            # Save the modified document
            docx_doc.save(output_path)

        file_size = output_path.stat().st_size
        message = f"✓ [{row_number:2d}] {recipient_name[:40]:40s} -> {filename} ({file_size:,} bytes)"
        return True, message

    except Exception as e:
        recipient_name = str(row.get('Recipient Name', 'Unknown'))[:40]
        message = f"✗ [{row_number:2d}] {recipient_name:40s} -> Error: {str(e)}"
        # Print detailed error for debugging
        import traceback
        print(f"\nDetailed error for {recipient_name}:")
        traceback.print_exc()
        return False, message


def main():
    """Main script execution."""
    print("=" * 80)
    print("RIR Cover Letter Generator - By Region")
    print("=" * 80)

    # Setup paths
    project_root = Path(__file__).parent.parent
    excel_path = project_root / "docs" / "RIR 2026" / "RIR Cover Letter" / "CORTAP Package 4 - Reviews - SM Updated 123025.xlsx"
    template_region1 = project_root / "app" / "templates" / "rir-cover-letter-region1.docx"
    template_region3 = project_root / "app" / "templates" / "rir-cover-letter-region3.docx"

    # Create timestamped output directories
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_base = project_root / "output" / "cover-letters-fy2026" / timestamp
    output_region1 = output_base / "region1"
    output_region3 = output_base / "region3"

    # Verify Excel file exists
    if not excel_path.exists():
        print(f"\n✗ Error: Excel file not found: {excel_path}")
        sys.exit(1)

    # Verify templates exist
    if not template_region1.exists():
        print(f"\n✗ Error: Region 1 template file not found: {template_region1}")
        sys.exit(1)

    if not template_region3.exists():
        print(f"\n✗ Error: Region 3 template file not found: {template_region3}")
        sys.exit(1)

    # Create output directories
    output_region1.mkdir(parents=True, exist_ok=True)
    output_region3.mkdir(parents=True, exist_ok=True)

    print(f"\nInput file: {excel_path.name}")
    print(f"Template Region 1: {template_region1.name}")
    print(f"Template Region 3: {template_region3.name}")
    print(f"Output directories:")
    print(f"  Region 1: {output_region1}")
    print(f"  Region 3: {output_region3}")

    # Read Excel file
    print(f"\nReading Excel file...")
    sheet_name = "RIR Cover Letter 2026 Info"
    try:
        # Read from the RIR Cover Letter 2026 Info sheet
        df = pd.read_excel(excel_path, sheet_name=sheet_name, header=0)

        print(f"  Sheet: {sheet_name}")
        print(f"  Found {len(df)} recipients")

    except Exception as e:
        print(f"\n✗ Error reading Excel file: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Split by region
    region1_df = df[df['Region'] == 'REGION I'].reset_index(drop=True)
    region3_df = df[df['Region'] == 'REGION III'].reset_index(drop=True)

    print(f"\n  Region I: {len(region1_df)} recipients")
    print(f"  Region III: {len(region3_df)} recipients")

    # Generate Region 1 documents
    print(f"\n{'=' * 80}")
    print(f"Generating {len(region1_df)} Region 1 cover letters...")
    print(f"{'=' * 80}\n")

    region1_results = []
    region1_messages = []

    for idx, row in region1_df.iterrows():
        success, message = generate_cover_letter_from_row(
            template_region1,
            row,
            output_region1,
            idx + 1
        )
        region1_results.append(success)
        region1_messages.append(message)
        print(message)

    # Generate Region 3 documents
    print(f"\n{'=' * 80}")
    print(f"Generating {len(region3_df)} Region 3 cover letters...")
    print(f"{'=' * 80}\n")

    region3_results = []
    region3_messages = []

    for idx, row in region3_df.iterrows():
        success, message = generate_cover_letter_from_row(
            template_region3,
            row,
            output_region3,
            idx + 1
        )
        region3_results.append(success)
        region3_messages.append(message)
        print(message)

    # Summary
    print(f"\n{'=' * 80}")
    print(f"Generation Summary")
    print(f"{'=' * 80}")
    print(f"Region I:")
    print(f"  Total: {len(region1_df)}")
    print(f"  Successful: {sum(region1_results)}")
    print(f"  Failed: {len(region1_results) - sum(region1_results)}")

    print(f"\nRegion III:")
    print(f"  Total: {len(region3_df)}")
    print(f"  Successful: {sum(region3_results)}")
    print(f"  Failed: {len(region3_results) - sum(region3_results)}")

    print(f"\nOverall:")
    print(f"  Total: {len(df)}")
    print(f"  Successful: {sum(region1_results) + sum(region3_results)}")
    print(f"  Failed: {(len(region1_results) - sum(region1_results)) + (len(region3_results) - sum(region3_results))}")

    if all(region1_results) and all(region3_results):
        print(f"\n✓ All {len(df)} cover letter documents generated successfully!")
        print(f"\nOutput directories:")
        print(f"  Region 1: {output_region1}")
        print(f"  Region 3: {output_region3}")
    else:
        print(f"\n⚠ Some documents failed to generate:")
        for success, message in zip(region1_results + region3_results, region1_messages + region3_messages):
            if not success:
                print(f"  {message}")


if __name__ == "__main__":
    main()
