#!/usr/bin/env python3
"""
Convert markdown documentation to Word document with proper tables.

Usage:
    python scripts/convert_md_to_word.py docs/draft-report-data-source-mapping.md output/draft-report-data-source-mapping.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style the hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def parse_markdown_table(lines):
    """Parse a markdown table into rows and columns."""
    # Remove empty lines
    lines = [line for line in lines if line.strip()]

    if len(lines) < 2:
        return None

    # Parse header
    header = [cell.strip() for cell in lines[0].split('|')[1:-1]]

    # Skip separator line (line 1)

    # Parse data rows
    rows = []
    for line in lines[2:]:
        if line.strip().startswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:  # Skip empty rows
                rows.append(cells)

    return {'header': header, 'rows': rows}


def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document."""
    if not table_data or not table_data['rows']:
        return

    # Create table
    table = doc.add_table(rows=1 + len(table_data['rows']), cols=len(table_data['header']))
    table.style = 'Light Grid Accent 1'

    # Add header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(table_data['header']):
        cell = header_cells[i]
        cell.text = header_text
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # Shade header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9E2F3')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Add data rows
    for row_idx, row_data in enumerate(table_data['rows'], start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row_cells):
                cell = row_cells[col_idx]
                # Handle emojis and formatting
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def process_inline_formatting(paragraph, text):
    """Process inline markdown formatting like bold, italic, code."""
    # Pattern for **bold**, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Regular text
            if part:
                paragraph.add_run(part)


def convert_markdown_to_word(md_file, output_file):
    """Convert markdown file to Word document."""
    # Read markdown file
    md_path = Path(md_file)
    if not md_path.exists():
        print(f"Error: File not found: {md_file}")
        return False

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process markdown line by line
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                # Add light gray background
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F0F0F0')
                p._element.get_or_add_pPr().append(shading_elm)

                in_code_block = False
                code_lines = []
            else:
                # Start of code block
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle tables
        if line.startswith('|') and not in_table:
            # Start of table
            in_table = True
            table_lines = [line]
            i += 1
            continue

        if in_table:
            if line.startswith('|'):
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                table_data = parse_markdown_table(table_lines)
                if table_data:
                    add_table_to_doc(doc, table_data)
                    doc.add_paragraph()  # Add space after table
                in_table = False
                table_lines = []
                # Don't increment i, process this line normally
                continue

        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.runs[0].font.size = Pt(20)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.size = Pt(16)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.runs[0].font.size = Pt(13)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.runs[0].font.size = Pt(11)
            heading.runs[0].font.bold = True

        # Handle horizontal rules
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(200, 200, 200)

        # Handle bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            # Check indent level
            indent_level = 0
            if line.startswith('  - ') or line.startswith('  * '):
                indent_level = 1
                text = line[4:]
            elif line.startswith('    - ') or line.startswith('    * '):
                indent_level = 2
                text = line[6:]
            else:
                text = line[2:]

            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * indent_level)

        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text)

        # Handle blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = 'Intense Quote'

        # Handle checkboxes
        elif line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            checked = 'x' in line.lower()[:6]
            text = line[6:] if len(line) > 6 else ''
            p = doc.add_paragraph()
            p.add_run('☑ ' if checked else '☐ ')
            process_inline_formatting(p, text)
            p.paragraph_format.left_indent = Inches(0.25)

        # Handle empty lines
        elif line.strip() == '':
            doc.add_paragraph()

        # Regular paragraphs
        else:
            if line.strip():
                p = doc.add_paragraph()
                process_inline_formatting(p, line)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        table_data = parse_markdown_table(table_lines)
        if table_data:
            add_table_to_doc(doc, table_data)

    # Save document
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"✅ Successfully converted {md_file} to {output_file}")
    return True


def main():
    if len(sys.argv) != 3:
        print("Usage: python scripts/convert_md_to_word.py <input.md> <output.docx>")
        print("Example: python scripts/convert_md_to_word.py docs/draft-report-data-source-mapping.md output/draft-report-data-source-mapping.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    success = convert_markdown_to_word(input_file, output_file)

    if not success:
        sys.exit(1)


if __name__ == '__main__':
    main()
