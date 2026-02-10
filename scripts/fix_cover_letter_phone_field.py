#!/usr/bin/env python3
"""
Fix phone field in RIR cover letter templates.

The last paragraph incorrectly uses {{ ro_rep_phone }} for both the RO rep
and the lead reviewer. The second instance should be {{ lead_phone }}.

Usage:
    python scripts/fix_cover_letter_phone_field.py
"""

from pathlib import Path
from docx import Document


def fix_phone_field_in_paragraph(paragraph):
    """
    Fix the phone field in a paragraph.

    The paragraph should have:
    "...contact {{ ro_rep_title_name }} by phone at {{ ro_rep_phone }} or {{ ro_rep_email }},
     or {{ lead_title_name }} of {{ contractor_name }} by phone at {{ lead_phone }} or {{ lead_email }}."

    But currently has {{ ro_rep_phone }} in both places.
    """
    text = paragraph.text

    # Check if this is the problematic paragraph
    # It should contain both ro_rep_phone and lead_email references
    if '{{ ro_rep_phone }}' in text and '{{ lead_email }}' in text:
        print(f"  Found paragraph with phone field issue")
        print(f"  Original: {text[:150]}...")

        # Count occurrences of ro_rep_phone
        count = text.count('{{ ro_rep_phone }}')
        if count == 2:
            # Find the second occurrence and replace it with lead_phone
            # The pattern is: "...by phone at {{ ro_rep_phone }} or {{ lead_email }}"
            # We want to replace the ro_rep_phone that comes before lead_email

            # Strategy: Replace the text in runs
            # We need to find the second occurrence of {{ ro_rep_phone }}
            found_first = False

            for run in paragraph.runs:
                if '{{ ro_rep_phone }}' in run.text:
                    if not found_first:
                        # This is the first occurrence, leave it alone
                        found_first = True
                        print(f"  Keeping first occurrence: '{run.text}'")
                    else:
                        # This is the second occurrence, replace it
                        print(f"  Replacing second occurrence: '{run.text}' -> ", end='')
                        run.text = run.text.replace('{{ ro_rep_phone }}', '{{ lead_phone }}')
                        print(f"'{run.text}'")
                        return True

            # If we didn't find it in separate runs, it might span runs or be in a single run
            # Try a different approach: find the run containing both ro_rep_phone and lead_email
            for run in paragraph.runs:
                if '{{ ro_rep_phone }}' in run.text and '{{ lead_email }}' in run.text:
                    # Replace only the second occurrence in this run
                    parts = run.text.split('{{ ro_rep_phone }}')
                    if len(parts) == 3:  # Two occurrences
                        # Reconstruct: first part + {{ ro_rep_phone }} + middle + {{ lead_phone }} + last part
                        run.text = parts[0] + '{{ ro_rep_phone }}' + parts[1] + '{{ lead_phone }}' + parts[2]
                        print(f"  Fixed in single run")
                        return True
                    elif len(parts) == 2:  # One occurrence, but we need context
                        # Check if this is the second one by looking for lead_email
                        if '{{ lead_email }}' in parts[1]:
                            run.text = parts[0] + '{{ lead_phone }}' + parts[1]
                            print(f"  Fixed single occurrence near lead_email")
                            return True

    return False


def fix_template(template_path: Path) -> bool:
    """
    Fix the phone field in a cover letter template.

    Args:
        template_path: Path to the template file

    Returns:
        True if successful, False otherwise
    """
    try:
        print(f"\nProcessing: {template_path.name}")

        # Open the template
        doc = Document(template_path)

        # Search all paragraphs for the problematic text
        fixed = False
        for paragraph in doc.paragraphs:
            if fix_phone_field_in_paragraph(paragraph):
                fixed = True
                break

        # Also check tables
        if not fixed:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if fix_phone_field_in_paragraph(paragraph):
                                fixed = True
                                break
                        if fixed:
                            break
                    if fixed:
                        break
                if fixed:
                    break

        if not fixed:
            print(f"  ⚠ Could not find the problematic paragraph")
            return False

        # Save the modified template
        doc.save(template_path)
        print(f"  ✓ Template fixed and saved")

        return True

    except Exception as e:
        print(f"  ✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Main script execution."""
    print("=" * 80)
    print("Fix Phone Field in RIR Cover Letter Templates")
    print("=" * 80)

    # Setup paths
    project_root = Path(__file__).parent.parent
    template_region1 = project_root / "app" / "templates" / "rir-cover-letter-region1.docx"
    template_region3 = project_root / "app" / "templates" / "rir-cover-letter-region3.docx"

    # Process both templates
    results = []

    if template_region1.exists():
        results.append(fix_template(template_region1))
    else:
        print(f"\n✗ Template not found: {template_region1}")
        results.append(False)

    if template_region3.exists():
        results.append(fix_template(template_region3))
    else:
        print(f"\n✗ Template not found: {template_region3}")
        results.append(False)

    # Summary
    print("\n" + "=" * 80)
    print("Summary")
    print("=" * 80)

    if all(results):
        print("✓ Both templates fixed successfully!")
        print("\nThe paragraph now correctly shows:")
        print("  ...contact {{ ro_rep_title_name }} by phone at {{ ro_rep_phone }} or {{ ro_rep_email }},")
        print("  or {{ lead_title_name }} of {{ contractor_name }} by phone at {{ lead_phone }} or {{ lead_email }}.")
    else:
        print("✗ Some templates failed to update")

    return all(results)


if __name__ == "__main__":
    import sys
    success = main()
    sys.exit(0 if success else 1)
